import os
import re
import json
import tempfile
import openpyxl
import markdown
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from google import genai
from google.genai import types
from .gcs_storage import get_gcs_manager
from pycel import ExcelCompiler
from . import config

def get_gemini_client():
    """Initializes and returns a Gemini client."""
    project = os.getenv("PROJECT")
    location = os.getenv("LOCATION")
    if not all([project, location]):
        raise ValueError("Missing environment variables for Gemini API (PROJECT, LOCATION).")
    return genai.Client(vertexai=True, project=project, location=location)

def get_summary_files(company_name):
    """Returns a list of all PDF blobs from specified GCS subfolders for the summary."""
    gcs = get_gcs_manager()
    summary_blobs = []
    folders_to_scan = ["annual_returns", "annual_reports", "shareholding_pattern"]
    for folder in folders_to_scan:
        prefix = f"{config.GCS_COMPANIES_FOLDER}/{company_name}/{folder}/"
        blobs = gcs.bucket.list_blobs(prefix=prefix)
        summary_blobs.extend([blob for blob in blobs if blob.name.endswith('.pdf')])
    return summary_blobs

def get_shareholding_files(company_name):
    """Returns a list of PDF blobs from the Shareholding Pattern GCS folder."""
    gcs = get_gcs_manager()
    prefix = f"{config.GCS_COMPANIES_FOLDER}/{company_name}/shareholding_pattern/"
    blobs = gcs.bucket.list_blobs(prefix=prefix)
    return [blob for blob in blobs if blob.name.endswith('.pdf')]

def get_spreadsheet_file(company_name):
    """
    Finds and returns the path to the downloaded spreadsheet file from GCS,
    with validation.
    """
    gcs = get_gcs_manager()
    prefix = f"{config.GCS_COMPANIES_FOLDER}/{company_name}/spreadsheet/"
    blobs = list(gcs.bucket.list_blobs(prefix=prefix))
    
    if not blobs:
        print(f"No files found in GCS at prefix: {prefix}")
        return None

    for blob in blobs:
        # We are specifically looking for .xlsx files as .xls is not supported by openpyxl
        if blob.name.endswith('.xlsx'):
            print(f"📥 Downloading spreadsheet: {blob.name} (Size: {blob.size} bytes)")
            
            # Create a temporary file to download to
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
            blob.download_to_filename(temp_file.name)
            
            # --- File Validation ---
            try:
                # Try to open the file with openpyxl to ensure it's a valid xlsx (zip) file
                openpyxl.load_workbook(temp_file.name)
                print(f"✅ Successfully validated spreadsheet: {temp_file.name}")
                return temp_file.name
            except Exception as e:
                print(f"❌ ERROR: The file '{blob.name}' is not a valid .xlsx file. Error: {e}")
                # Clean up the invalid downloaded file
                os.unlink(temp_file.name)
                # Continue to the next file if there are others
                continue
                
    print(f"🟡 No valid .xlsx spreadsheet found for {company_name}.")
    return None

def get_prompt(prompt_name):
    """
    Downloads a specific prompt text file from GCS.
    """
    gcs = get_gcs_manager()
    prompt_filename = f"{prompt_name}.txt"
    gcs_path = f"{config.GCS_PROMPTS_FOLDER}/{prompt_filename}"
    
    print(f"📥 Downloading prompt '{prompt_name}' from GCS path: {gcs_path}")
    
    try:
        blob = gcs.bucket.blob(gcs_path)
        if not blob.exists():
            raise FileNotFoundError(f"Prompt file not found in GCS at '{gcs_path}'")
        
        prompt_content = blob.download_as_text()
        return prompt_content
        
    except Exception as e:
        print(f"❌ Failed to download prompt '{prompt_name}' from GCS: {e}")
        raise

def read_prompt(prompt_filename):
    """Reads a prompt file from the local prompts directory."""
    prompt_path = os.path.join(config.PROMPTS_DIR, prompt_filename)
    try:
        with open(prompt_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        print(f"Error: Prompt file not found at {prompt_path}")
        return ""
    except Exception as e:
        print(f"Error reading prompt file: {e}")
        return ""

def read_financial_data(excel_path, header_row=2, data_start_row=3):
    """
    Reads financial data from the specified Excel file, saves it as a Markdown table,
    and returns the table as a string.
    """
    try:
        workbook = openpyxl.load_workbook(excel_path, data_only=True)
        sheet = workbook[config.FINANCIALS_SHEET_NAME]
        
        headers = [str(cell.value) if cell.value is not None else "" for cell in sheet[config.HEADER_ROW]]
        
        data_rows = []
        num_columns = len(headers)
        for row in sheet.iter_rows(min_row=config.DATA_START_ROW, values_only=True):
            if any(row):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if len(row_data) < num_columns:
                    row_data.extend([""] * (num_columns - len(row_data)))
                data_rows.append(row_data[:num_columns])
        
        markdown_table = "| " + " | ".join(headers) + " |\n"
        markdown_table += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for row in data_rows:
            markdown_table += "| " + " | ".join(row) + " |\n"

        # Define the output path and ensure the directory exists
        output_dir = config.OUTPUT_DATA_DIR
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, "financial_data.md")

        # Save the markdown table to the specified file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_table)
        print(f"Successfully saved financial data to: {output_path}")
            
        return markdown_table
        
    except FileNotFoundError:
        raise FileNotFoundError(f"Excel file not found at '{excel_path}'")
    except Exception as e:
        raise IOError(f"An error occurred while reading the Excel file: {e}")

def format_cell(cell):
    """
    Formats a cell's value, preserving percentage formatting.
    """
    if cell.value is None:
        return ""
    # Check if the format is percentage AND the value is a number
    if cell.number_format and '%' in cell.number_format and isinstance(cell.value, (int, float)):
        # Format the number as a percentage string with 2 decimal places
        return f"{cell.value * 100:.2f}%"
    return str(cell.value)

def format_value(value, number_format):
    """
    Formats a cell's value, preserving percentage formatting.
    """
    if value is None:
        return ""
    # Check if the format is percentage AND the value is a number
    if number_format and '%' in number_format and isinstance(value, (int, float)):
        # Format the number as a percentage string with 2 decimal places
        return f"{value * 100:.2f}%"
    return str(value)


def extract_financial_analysis_table(excel_path):
    """
    Extracts financial data from rows 2-34 of the 'Financials' sheet in an Excel file
    and returns it as a Markdown table.
    """
    try:
        # Load the workbook with openpyxl to get number formats
        workbook = openpyxl.load_workbook(excel_path, data_only=False)
        sheet_name = config.FINANCIALS_SHEET_NAME
        sheet = workbook[sheet_name]

        # Use ExcelCompiler to evaluate the formulas
        excel = ExcelCompiler(filename=excel_path)
        
        # Headers are in cells B2, C2, D2, E2
        headers = [
            excel.evaluate(f'{sheet_name}!B2'),
            excel.evaluate(f'{sheet_name}!C2'),
            excel.evaluate(f'{sheet_name}!D2'),
            excel.evaluate(f'{sheet_name}!E2')
        ]
        headers = [str(h) if h is not None else "" for h in headers]

        data_rows = []
        # Data is in the range B4:E28
        for row_index in range(4, 29):
            row_data = []
            for col_index in range(2, 6): # Columns B, C, D, E
                cell_address = f"{openpyxl.utils.get_column_letter(col_index)}{row_index}"
                
                # Get the calculated value from PyCel
                calculated_value = excel.evaluate(f'{sheet_name}!{cell_address}')
                
                # Get the number format from openpyxl
                number_format = sheet[cell_address].number_format
                
                # Format the value
                formatted_value = format_value(calculated_value, number_format)
                row_data.append(formatted_value)
            data_rows.append(row_data)

        # Build the Markdown table
        markdown_table = "| " + " | ".join(headers) + " |\n"
        markdown_table += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for row in data_rows:
            markdown_table += "| " + " | ".join(row) + " |\n"
        
        context_string = f"### Data from Excel Tab: {sheet_name}\n\n"
        context_string += markdown_table
        
        # output_dir = "data/output"
        # os.makedirs(output_dir, exist_ok=True)
        # output_path = os.path.join(output_dir, "financial_data.md")

        # # Save the markdown table to the specified file
        # with open(output_path, 'w', encoding='utf-8') as f:
        #     f.write(context_string)
        # print(f"Successfully saved financial data to: {output_path}")
        
        return context_string

    except FileNotFoundError:
        print(f"Error: The file was not found at '{file_path}'")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")


def convert_md_to_docx(md_content, docx_path):
    """
    Converts Markdown content to a DOCX file, handling headers, paragraphs,
    tables, and lists, ensuring robust HTML parsing.
    """
    # 1. Convert Markdown to HTML and wrap it in a proper structure
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    # *** CRITICAL FIX: Wrap in <html><body> tags for reliable BeautifulSoup parsing ***
    wrapped_html = f"<html><body>{html_content}</body></html>" 
    
    # Save the intermediate HTML file (Good for debugging)
    html_path = docx_path.replace(".docx", ".html")
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(wrapped_html)
    print(f"Successfully saved intermediate HTML to: {html_path}")

    # 2. Parse HTML
    soup = BeautifulSoup(wrapped_html, 'html.parser')
    doc = Document()
    
    # Define the tags we want to process
    content_tags = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table', 'ol', 'ul']

    # Iterate over the direct children of the <body> tag to preserve order and avoid duplication
    # We use 'body' as the container because we wrapped the content in it.
    body = soup.find('body')
    if body:
        # Use recursive=False to only find top-level elements, which prevents duplicating content
        for element in body.find_all(content_tags, recursive=False):
            
            # --- Handle Headers ---
            if element.name.startswith('h'):
                # Extract the header level number and use it to set the docx level
                try:
                    level = int(element.name[1])
                    doc.add_heading(element.text.strip(), level=level)
                except ValueError:
                    doc.add_paragraph(element.text.strip()) 
                    
            # --- Handle Paragraphs ---
            elif element.name == 'p':
                doc.add_paragraph(element.text.strip())
                
            # --- Handle Tables ---
            elif element.name == 'table':
                table_data = []
                # First, find the header row (th) to determine the number of columns
                header_row = element.find('thead').find('tr') if element.find('thead') else element.find('tr')
                if header_row:
                    num_cols = len(header_row.find_all(['th', 'td']))
                    # Gather all row data
                    for row in element.find_all('tr'):
                        row_data = [cell.text.strip() for cell in row.find_all(['th', 'td'])]
                        if row_data:
                            # Pad/truncate row data to match column count
                            row_data = (row_data + [''] * num_cols)[:num_cols] 
                            table_data.append(row_data)

                if table_data:
                    num_rows = len(table_data)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            table.cell(i, j).text = cell_text

            # --- Handle Lists ---
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text.strip(), style='List Bullet')
                    
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text.strip(), style='List Number')
    else:
        print("Error: Could not find the <body> element in the parsed HTML.")


    doc.save(docx_path)
    print(f"Successfully converted report to DOCX at: {docx_path}")

# def convert_md_to_docx(md_content, docx_path):
#     """
#     Converts Markdown content to a DOCX file with a table of contents,
#     and saves the intermediate HTML.
#     """
#     html = markdown.markdown(md_content, extensions=['tables'])
    
#     # Save the intermediate HTML file
#     html_path = docx_path.replace(".docx", ".html")
#     with open(html_path, 'w', encoding='utf-8') as f:
#         f.write(html)
#     print(f"Successfully saved intermediate HTML to: {html_path}")

#     soup = BeautifulSoup(html, 'html.parser')
    
#     doc = Document()
    
#     # Add the rest of the content
#     for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table', 'ol', 'ul']):
#         if element.name.startswith('h'):
#             doc.add_heading(element.text, level=1)
#         elif element.name == 'p':
#             doc.add_paragraph(element.text)
#         elif element.name == 'table':
#             table_data = []
#             for row in element.find_all('tr'):
#                 row_data = [cell.text for cell in row.find_all(['th', 'td'])]
#                 table_data.append(row_data)
            
#             if table_data:
#                 num_rows = len(table_data)
#                 num_cols = len(table_data[0])
#                 table = doc.add_table(rows=num_rows, cols=num_cols)
#                 table.style = 'Table Grid'
#                 for i, row_data in enumerate(table_data):
#                     for j, cell_text in enumerate(row_data):
#                         table.cell(i, j).text = cell_text
#         elif element.name == 'ul':
#             for li in element.find_all('li'):
#                 doc.add_paragraph(li.text, style='List Bullet')
#         elif element.name == 'ol':
#             for li in element.find_all('li'):
#                 doc.add_paragraph(li.text, style='List Number')

#     doc.save(docx_path)
#     print(f"Successfully converted report to DOCX at: {docx_path}")

def format_markdown_headings(input_file_path, output_file_path):
    """
    Reads a markdown file, formats its headings to ensure a clear hierarchy,
    and writes the output to a new file.

    Main sections (e.g., Section 1, Section 2) will be treated as Heading 1 (#).
    All other headings will be adjusted to a proper nested hierarchy.
    """
    try:
        with open(input_file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # Regular expressions to match headings
        section_heading_re = re.compile(r'^(#+)\s*Section\s*\d+:')
        other_heading_re = re.compile(r'^(#+)\s*(?!Section)')

        formatted_lines = []
        for line in lines:
            line_stripped = line.strip()

            # Check for "Section" headings and ensure they are exactly Heading 1 (#)
            section_match = section_heading_re.match(line_stripped)
            if section_match:
                content = section_match.group(0).replace(section_match.group(1), '#')
                formatted_lines.append(content + line_stripped[len(section_match.group(0)):].strip() + '\n\n')
                continue

            # Check for other headings and standardize them to Heading 2 (##)
            other_heading_match = other_heading_re.match(line_stripped)
            if other_heading_match:
                # Standardize all other headings to ##
                content = '##' + line_stripped.lstrip('#').strip()
                formatted_lines.append(content + '\n\n')
                continue
            
            # For all other lines, keep them as they are
            formatted_lines.append(line)

        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.writelines(formatted_lines)
            
        print(f"Successfully formatted '{input_file_path}' and saved to '{output_file_path}'.")

    except FileNotFoundError:
        print(f"Error: The file '{input_file_path}' was not found.")
    except Exception as e:
        print(f"An error occurred: {e}")

def clean_ai_output(text):
    """Removes leading/trailing Markdown code fences and leading headings."""
    text = text.strip()
    
    # Remove markdown code fences if present
    if text.startswith("```"):
        lines = text.split('\n')
        # Remove first line if it's a fence found at start
        if lines[0].startswith("```"):
            lines = lines[1:]
        # Remove last line if it's a fence
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        text = '\n'.join(lines).strip()

    lines = text.split('\n')
    if lines and lines[0].strip().startswith('#'):
        # Heuristic: if first line is a heading, it might be a title we want to strip
        # But be careful not to strip useful headers if they are part of content
        # For now, let's keep the user's original logic but aware of the fence removal above
        return '\n'.join(lines[1:]).lstrip()
    return text

def format_markdown_with_llm(unformatted_markdown):
    """
    Uses an LLM to reformat a Markdown document.
    """
    client = get_gemini_client()

    prompt = f"""Here is the Markdown document that needs to be cleaned up:
{unformatted_markdown}

    Reformat the entire document based on the following rules:

1.  All headings must use the ATX style (e.g., #, ##).
2.  Ensure there is a single space after the # and before the heading text.
3.  Use a single blank line before and after each heading and list.
4.  Standardize unordered lists to use a single hyphen (-) and a space.
5.  Standardize ordered lists to use a single 1. followed by a space.
6.  For bold text, use **double asterisks**. For italics, use *single asterisks*.
7.  Ensure that the top-level heading for each section is h1, and understand the hierarchy underneath to have appropriate headings.
8.  Do not include any pre-amble like "here is the formated markdown based on the rules".

Example:

Input:

Some Text
## Heading 2
- item1
*bold*
_italics_

Output:

Some Text

## Heading 2

- item1

**bold**

*italics*"""

    system_instruction = """You are a Markdown formatting expert. Your task is to reformat a given Markdown document according to a specific set of rules. You will receive the Markdown content, and you must apply the formatting rules precisely."""

    model = config.GEMINI_PRO_MODEL
    contents = [types.Content(role="user", parts=[types.Part.from_text(text=prompt)])]

    generate_content_config = types.GenerateContentConfig(
        temperature=1,
        top_p=0.95,
        max_output_tokens=65535,
        system_instruction=[types.Part.from_text(text=system_instruction)],
    )

    response = client.models.generate_content(
        model=model,
        contents=contents,
        config=generate_content_config
    )
    
    return response.text

def call_gemini_model(
    prompt_text: str,
    parts: list = None,
    model_name: str = None,
    tools: list = None,
    thinking_enabled: bool = False,
    safety_settings: list = None,
    temperature: float = None,
    max_output_tokens: int = 65535,
    client: genai.Client = None
):
    """
    Common utility to call Gemini model with standardized configuration.
    
    Args:
        prompt_text: The text prompt.
        parts: Optional list of existing parts (if provided, prompt_text is appended/prepended or verified). 
               If None, a text part is created from prompt_text.
        model_name: Model ID to use (default: configuration.GEMINI_FLASH_MODEL).
        tools: List of tools (e.g., [types.Tool(google_search=types.GoogleSearch())]).
        thinking_enabled: Whether to enable ThinkingConfig (Thinking Budget).
        safety_settings: Custom safety settings. If None, uses permissive defaults.
        temperature: Generation temperature.
        max_output_tokens: Max tokens.
        client: Pre-initialized client. If None, creates a new one.
        
    Returns:
        str: The generated text response.
    """
    if client is None:
        client = get_gemini_client()
        
    if model_name is None:
        model_name = config.GEMINI_FLASH_MODEL

    # Construct content parts
    if parts is None:
        parts = [types.Part.from_text(text=prompt_text)]
    else:
        # If parts list exists, we assume the caller handled the prompt text construction 
        # or we can check if prompt_text needs to be added. 
        # For simplicity, if parts is passed, we assume it's complete, 
        # but if prompt_text is ALSO passed and not in parts, might be ambiguous.
        # Let's assume if parts is given, prompt_text is largely ignored or already in parts.
        # BUT, to be safe:
        if prompt_text and not any(p.text == prompt_text for p in parts if hasattr(p, 'text')):
             # Prepend or append? Usually prompt text is instructions.
             parts.insert(0, types.Part.from_text(text=prompt_text))

    contents = [types.Content(role="user", parts=parts)]

    # Safety Settings (Default: OFF for all if not provided)
    if safety_settings is None:
        safety_settings = [
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="OFF"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="OFF"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="OFF"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="OFF")
        ]

    # Thinking Config
    thinking_config = None
    if thinking_enabled:
        thinking_config = types.ThinkingConfig(thinking_budget=-1)
    
    # Generate Config
    gen_config_args = {
        "max_output_tokens": max_output_tokens,
        "safety_settings": safety_settings,
        "tools": tools,
        "thinking_config": thinking_config
    }
    
    if temperature is not None:
        gen_config_args["temperature"] = temperature
    elif not thinking_enabled:
        # Only set default temperature if thinking is disabled 
        # (Thinking models often handle temperature differently or mandate it)
        gen_config_args["temperature"] = config.TEMPERATURE # Default 0.1 or from env

    generate_content_config = types.GenerateContentConfig(**gen_config_args)

    try:
        response = client.models.generate_content(
            model=model_name,
            contents=contents,
            config=generate_content_config
        )
        return clean_ai_output(response.text)
    except Exception as e:
        print(f"❌ Error calling Gemini model: {e}")
        raise

# --- DOCX STYLING UTILITIES ---
from docx.oxml.ns import qn

def apply_font_to_run(run, font_name):
    """
    Helper function to robustly apply font name to a run.
    Uses XML override to ensure it sticks for all character types.
    """
    run.font.name = font_name
    
    # --- XML FORCE OVERRIDE ---
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    
    # Force the name on all specific font types (ASCII, High ANSI, Complex Script)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

def process_document(
    input_docx,
    output_docx,
    table_style="Grid Table 1 Light",
    font_name="Calibri",      # Default changed to Calibri as per requirement
    total_width=6.0,
    min_col_width=0.8,
    max_col_width=3.0
):
    """
    Standardizes fonts and table layouts in a DOCX file.
    """
    try:
        doc = Document(input_docx)

        # ---------------------------------------------------------
        # PART 1: Apply Font to Main Body Paragraphs
        # ---------------------------------------------------------
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                apply_font_to_run(run, font_name)

        # ---------------------------------------------------------
        # PART 2: Apply Font + Resizing to Tables
        # ---------------------------------------------------------
        for table in doc.tables:
            # 1️⃣ Apply Word table style
            try:
                table.style = table_style
            except KeyError:
                pass # Skip if style doesn't exist

            # 2️⃣ Measure content-aware column weights
            col_weights = []
            for col in table.columns:
                if not col.cells:
                    col_weights.append(1)
                    continue
                    
                header_len = len(col.cells[0].text.strip())
                body_cells = col.cells[1:] if len(col.cells) > 1 else col.cells
                body_len = max(
                    (len(cell.text.strip()) for cell in body_cells),
                    default=1
                )
                weight = max(header_len * 1.2, body_len)
                col_weights.append(max(weight, 1))

            total_weight = sum(col_weights)

            # 3️⃣ Convert weights → fixed widths
            widths = []
            for w in col_weights:
                width = total_width * (w / total_weight)
                width = max(min_col_width, min(width, max_col_width))
                widths.append(Inches(width))

            # 4️⃣ Disable Word autofit
            table.autofit = False
            table.allow_autofit = False

            # 5️⃣ Apply widths AND Font to all cells
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    # Apply Width
                    if i < len(widths):
                        cell.width = widths[i]
                    
                    # Apply Font to paragraphs inside the table cell
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            apply_font_to_run(run, font_name)

        doc.save(output_docx)
        print(f"✅ Document styled and saved: {output_docx} (Font: {font_name})")
        
    except Exception as e:
        print(f"⚠️ Warning: Failed to process/style document {input_docx}: {e}")
